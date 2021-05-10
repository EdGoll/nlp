import os
import stanza 
import docx2txt
from docx2python import docx2python
import docx
import numpy as np
import pandas as pd
import json

dir_docx="e:/workspace/python/FONDEF/MamAI/src/ejemplos_aleatorios"
archives = os.listdir(dir_docx)
str_token =':'
hallazgos_text = ''
impresion_text = ''
for file in archives :
    print(file)
    doc = docx.Document(dir_docx+"\\"+file)
    all_paras = doc.paragraphs
    len(all_paras)
    for para in all_paras:        
        #print("-------")
        #print(para.text)
        if para.text == 'Hallazgos:' :
            str_token ='Hallazgos:'
            continue

        if str_token == 'Hallazgos:' and para.text != '':
            #print("Hallazgos "+para.text)  
            if 'Impresión:' in para.text:                
                str_token ='Impresión:'
                continue
            else:
                print("Hallazgos "+para.text)  
                hallazgos_text += ''.join(para.text)
        if str_token == 'Impresión:' and para.text != '' :
            #if para.text not in 'BIRADS':
            if para.text.find("BIRADS") == -1:
                print("Impresión "+para.text)
                impresion_text += ''.join(para.text)
            else :
                break
    out = dict()
    out['impresiones'] = impresion_text
    out['hallazgos'] = hallazgos_text

    print(out)

    break


def findHallazgos(all_paras) :
    for para in all_paras:        
        if 'Hallazgos:' in para.text:
            print("Hallazgos "+para.text)
            continue
        if 'Atentamente' in para.text:
            break

def findImpresion(all_paras) : 
    for para in all_paras:
        #print(para.text)
        if 'Impresión:' in para.text:
            print("Impresión "+para.text)
            continue
        if 'Atentamente' in para.text:
            break

file = "7297718.docx"
#doc = docx.Document(dir_docx+"\\"+file)
all_paras = doc.paragraphs
#print(len(all_paras))
#findHallazgos(all_paras)
#findImpresion(all_paras)


